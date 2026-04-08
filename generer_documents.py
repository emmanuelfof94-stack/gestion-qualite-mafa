#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script de génération des documents Gestion Qualité
- QR Codes des agents
- Règlement Intérieur (Word)
- Fichier de pointage (Excel)
- Registre des agents (Excel)
"""

import json
import os
import qrcode
from datetime import datetime, date
from pathlib import Path

BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
QR_DIR = BASE_DIR / "qrcodes"
EXPORT_DIR = BASE_DIR / "exports"

# ─────────────────────────────────────────────
# 1. GÉNÉRATION DES QR CODES
# ─────────────────────────────────────────────
def generer_qrcodes():
    with open(DATA_DIR / "agents.json", encoding="utf-8") as f:
        agents = json.load(f)

    for agent in agents:
        data = f"BADGE|{agent['id']}|{agent['nom']}|{agent['prenom']}"
        qr = qrcode.QRCode(version=2, box_size=10, border=4,
                           error_correction=qrcode.constants.ERROR_CORRECT_H)
        qr.add_data(data)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        path = QR_DIR / f"QR_{agent['id']}_{agent['nom']}.png"
        img.save(str(path))
        print(f"  ✓ QR Code généré : {path.name}")

# ─────────────────────────────────────────────
# 2. RÈGLEMENT INTÉRIEUR (WORD)
# ─────────────────────────────────────────────
def generer_reglement_word():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # Marges
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── En-tête ──
    header = doc.add_heading("", 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("RÈGLEMENT INTÉRIEUR")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x76)

    sous_titre = doc.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sous_titre.add_run("COLIBRI TECHNOLOGIES — Service Gestion Qualité")
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run(f"Date d'entrée en vigueur : {date.today().strftime('%d/%m/%Y')}").font.size = Pt(10)

    doc.add_paragraph()

    # ── Introduction ──
    intro = doc.add_paragraph()
    intro.add_run(
        "Le présent règlement intérieur s'applique à l'ensemble du personnel affecté au service. "
        "Son respect est obligatoire et conditionne le bon fonctionnement collectif de notre environnement de travail. "
        "Tout manquement pourra faire l'objet d'une mesure disciplinaire."
    ).font.size = Pt(11)

    doc.add_paragraph()

    BLEU   = RGBColor(0x1A, 0x56, 0x76)
    ROUGE  = RGBColor(0xC0, 0x39, 0x2B)
    VERT   = RGBColor(0x1E, 0x84, 0x49)
    VIOLET = RGBColor(0x7D, 0x3C, 0x98)

    def ajouter_chapitre(titre, couleur=BLEU):
        p = doc.add_paragraph()
        run = p.add_run(titre.upper())
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = couleur
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)

    def ajouter_article(num, titre, contenu_liste, couleur=BLEU, intro=None):
        h = doc.add_heading("", level=2)
        h.clear()
        r_num = h.add_run(f"{num} — ")
        r_num.font.color.rgb = couleur
        r_num.font.size = Pt(12)
        r_num.font.bold = True
        r_titre = h.add_run(titre)
        r_titre.font.color.rgb = couleur
        r_titre.font.size = Pt(12)
        r_titre.font.bold = True
        if intro:
            p = doc.add_paragraph()
            p.add_run(intro).font.size = Pt(11)
        for ligne in contenu_liste:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(ligne).font.size = Pt(11)
        doc.add_paragraph()

    # ─── Chapitre I ───
    ajouter_chapitre("Chapitre I — Organisation Technique du Travail")

    ajouter_article("Art. 1", "Duree du Travail", [
        "Duree legale : 8 heures par jour, soit 40 heures par semaine (Loi N°2015-532 du 20/07/2015).",
        "Personnel administratif : Lundi au Vendredi — 09h00-12h30 et 14h30-17h30.",
        "Travailleurs sur terrain : Lundi au Vendredi — 08h00-12h00 et 13h30-16h30.",
        "Toute heure supplementaire doit faire l'objet d'une autorisation prealable du responsable.",
    ])

    ajouter_article("Art. 2", "Absences, Retards et Permissions Exceptionnelles", [
        "Toute absence doit etre prealablement autorisee par le Superieur hierarchique.",
        "Le travailleur empeche de se presenter doit immediatement prevenir son Superieur en precisant le motif.",
        "Les retards ou absences injustifies feront l'objet de sanctions.",
        "Absence > 72h sans notification = abandon de poste pouvant entrainer la rupture du contrat.",
        "Mariage du travailleur : 4 jours ouvrables.",
        "Mariage d'un enfant, frere ou soeur : 2 jours ouvrables.",
        "Deces du conjoint / enfant / pere / mere : 5 jours ouvrables.",
        "Deces d'un frere ou d'une soeur : 2 jours ouvrables.",
        "Deces d'un beau-pere ou belle-mere : 2 jours ouvrables.",
        "Naissance d'un enfant : 2 jours ouvrables.",
        "Bapteme d'un enfant : 1 jour ouvrable.",
        "Demenagement : 1 jour ouvrable.",
        "Toute permission doit etre autorisee par ecrit au plus tard 7 jours ouvrables avant.",
        "Les pieces justificatives doivent etre presentees dans les 3 jours suivant l'evenement.",
    ])

    ajouter_article("Art. 3", "Absences pour Maladies", [
        "En cas de maladie, prevenir l'employeur dans un delai de 72 heures, sauf force majeure.",
        "Le travailleur est tenu d'accepter la contre-visite du medecin d'entreprise.",
        "Tout refus de contre-visite peut entrainer le licenciement.",
    ])

    ajouter_article("Art. 4", "Accidents de Travail", [
        "Tout accident, meme leger, doit etre declare dans les 24 heures au Superieur hierarchique.",
        "L'information doit etre transmise aux RH et au Service Hygiene et Securite.",
        "Cette obligation s'applique egalement lors des deplacements professionnels.",
    ])

    ajouter_article("Art. 5", "Conges Annuels", [
        "Le droit aux conges est acquis apres 1 annee de service effectif.",
        "Le calendrier est etabli en accord avec le Superieur, selon les imperatifs de l'entreprise.",
        "La demande de conge doit etre soumise au moins 2 semaines avant la date de depart prevue.",
        "Chaque salarie est informe de ses dates de conge au moins 15 jours a l'avance.",
        "Le rappel en conge ne peut intervenir que pour les necessites de service.",
    ])

    # ─── Chapitre II ───
    ajouter_chapitre("Chapitre II — Discipline Generale dans la Societe")

    ajouter_article("Art. 6", "Obligations du Personnel", [
        "Entretenir des rapports de respect et de courtoisie avec superieurs, collegues et tiers.",
        "Exercer ses fonctions avec conscience, honnetete et devouement.",
        "Observer une discretion absolue dans l'execution de sa tache.",
        "Se presenter au poste de travail correctement et proprement vetu.",
    ])

    ajouter_article("Art. 7", "Interdictions", [
        "Exercer une activite concurrente ou nuisant a la bonne execution des services.",
        "Divulguer les renseignements detenus sur la societe.",
        "Entrer dans les locaux en etat d'ivresse ou sous substances illicites.",
        "Dormir pendant les heures de travail.",
        "S'adonner a toute occupation personnelle pendant les heures de travail.",
        "Introduire des marchandises pour les vendre ou les entreposer.",
        "Accepter des pots-de-vin ou accorder des avantages indus.",
        "Emporter sans autorisation des documents ou materiels de COLIBRI TECHNOLOGIES.",
        "Permettre l'acces aux equipements informatiques a des personnes etrangeres.",
        "Causer du desordre ou tenir des propos contraires aux bonnes moeurs.",
        "Faire pression sur un subordonne pour un travail contraire a l'objet social.",
        "Permettre l'utilisation de vehicules de l'entreprise a des personnes etrangeres.",
        "Emprunter un vehicule sans autorisation prealable ecrite.",
        "Adopter tout comportement raciste, xenophobe, sexiste ou discriminant.",
    ], couleur=ROUGE, intro="Il est formellement interdit a l'ensemble du personnel de COLIBRI TECHNOLOGIES :")

    ajouter_article("Art. 8", "Procedure Disciplinaire", [
        "Prealablement a toute sanction, le travailleur dispose de 72 heures pour s'expliquer.",
        "L'explication peut etre fournie par ecrit ou verbalement (article 17.5 du Code du travail).",
    ])

    ajouter_article("Art. 9", "Sanctions Disciplinaires", [
        "Niveau 1 — Avertissement ecrit.",
        "Niveau 2 — Mise a pied temporaire sans salaire (1 a 3 jours).",
        "Niveau 3 — Mise a pied temporaire sans salaire (4 a 8 jours).",
        "Niveau 4 — Licenciement.",
    ], couleur=ROUGE)

    ajouter_article("Art. 10", "Licenciement pour Faute Lourde", [
        "Incitation du personnel a la desobeissance.",
        "Etat d'ivresse ou influence de substances illicites.",
        "Infraction aux regles de securite ou violences physiques.",
        "Soustraction, meme temporaire, de documents.",
        "Deterioration volontaire d'un materiel.",
        "Insultes, menaces, voies de fait envers le personnel.",
        "Absence non motivee, repetee ou prolongee.",
        "Abandon de poste.",
        "Insubordination ou manque de respect caracterise.",
        "Mauvaise volonte persistante dans l'accomplissement de sa tache.",
        "Detournement de valeurs, objets ou fonds appartenant a l'entreprise.",
    ], couleur=ROUGE, intro="Le licenciement sans preavis ni indemnites pourra etre prononce notamment pour :")

    # ─── Chapitre III ───
    ajouter_chapitre("Chapitre III — Hygiene, Securite et Environnement", couleur=VERT)

    ajouter_article("Art. 11", "Generalites HSE", [
        "HYGIENE — Le personnel doit se presenter au travail en parfait etat de proprete corporelle et vestimentaire.",
        "HYGIENE — Les locaux et les lieux d'aisance doivent etre laisses propres apres usage.",
        "SECURITE — Chacun doit respecter et faire respecter les consignes de securite.",
        "SECURITE — Toute situation a risque doit etre signalee immediatement au superieur ou au responsable HSE.",
        "ENVIRONNEMENT — Le personnel doit identifier et reduire les impacts de ses activites sur l'environnement.",
        "ENVIRONNEMENT — Comportement responsable dans l'utilisation de l'energie, de l'eau et des ressources.",
        "ENVIRONNEMENT — Prevenir toute situation d'urgence (pollution, incendie...) pour l'environnement.",
    ], couleur=VERT)

    # ─── Chapitre IV ───
    ajouter_chapitre("Chapitre IV — Gestion des Missions", couleur=VIOLET)

    ajouter_article("Art. 12", "Procedure de Mission", [
        "PREPARATION — Completer le fichier Excel de demande de mission avec toutes les informations requises.",
        "VALIDATION — Soumettre la demande au Superieur hierarchique pour approbation.",
        "VALIDATION — Transmettre ensuite a la Direction Administrative et Financiere (DAF) pour traitement.",
        "FINANCEMENT — La DAF procede a la demande de financement aupres de MAFA Holding.",
        "FINANCEMENT — Le versement est effectue via la plateforme de paiement Julaya.",
        "POST-MISSION — Etablir un bilan detaille de toutes les depenses effectuees.",
        "POST-MISSION — En cas d'excedent budgetaire, restituer le montant en especes.",
        "POST-MISSION — En cas de depassement justifie, la DAF rembourse les frais supplementaires.",
    ], couleur=VIOLET)

    # ── Signatures ──
    doc.add_page_break()
    sig_titre = doc.add_paragraph()
    sig_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_titre.add_run("SIGNATURES").font.bold = True

    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Le Responsable Qualité"
    table.cell(0, 1).text = "L'Agent (Lu et approuvé)"
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    table.cell(0, 1).paragraphs[0].runs[0].bold = True
    table.cell(1, 0).text = "\n\nNom : ___________________"
    table.cell(1, 1).text = "\n\nNom : ___________________"
    table.cell(2, 0).text = "Signature : ______________"
    table.cell(2, 1).text = "Signature : ______________"

    out = EXPORT_DIR / "Reglement_Interieur.docx"
    doc.save(str(out))
    print(f"  ✓ Règlement Intérieur Word : {out.name}")

# ─────────────────────────────────────────────
# 3. FICHIER EXCEL — POINTAGE
# ─────────────────────────────────────────────
def generer_excel_pointage():
    from openpyxl import Workbook
    from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side,
                                  GradientFill)
    from openpyxl.utils import get_column_letter
    from openpyxl.formatting.rule import ColorScaleRule, CellIsRule, FormulaRule
    from openpyxl.styles.differential import DifferentialStyle

    with open(DATA_DIR / "agents.json", encoding="utf-8") as f:
        agents = json.load(f)

    wb = Workbook()

    # ── Couleurs ──
    BLEU_FONCE = "1A5676"
    BLEU_CLAIR = "D0E8F2"
    VERT       = "C6EFCE"
    ROUGE      = "FFC7CE"
    ORANGE     = "FFEB9C"
    GRIS       = "F2F2F2"

    def style_header(cell, bg=BLEU_FONCE, fg="FFFFFF", bold=True, size=11):
        cell.font = Font(bold=bold, color=fg, size=size)
        cell.fill = PatternFill("solid", fgColor=bg)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def border_all(ws, min_row, max_row, min_col, max_col):
        thin = Side(style="thin", color="AAAAAA")
        brd = Border(left=thin, right=thin, top=thin, bottom=thin)
        for row in ws.iter_rows(min_row=min_row, max_row=max_row,
                                 min_col=min_col, max_col=max_col):
            for cell in row:
                cell.border = brd

    # ════════════════════════════════════════
    # FEUILLE 1 — POINTAGE JOURNALIER
    # ════════════════════════════════════════
    ws1 = wb.active
    ws1.title = "Pointage Journalier"
    ws1.sheet_view.showGridLines = False

    # Titre
    ws1.merge_cells("A1:I1")
    ws1["A1"] = "FEUILLE DE POINTAGE JOURNALIER — SERVICE QUALITÉ"
    style_header(ws1["A1"], size=14)
    ws1.row_dimensions[1].height = 30

    # Date
    ws1.merge_cells("A2:I2")
    ws1["A2"] = f"Mois : {datetime.now().strftime('%B %Y').upper()}"
    ws1["A2"].font = Font(bold=True, size=11, color=BLEU_FONCE)
    ws1["A2"].alignment = Alignment(horizontal="center")

    # En-têtes colonnes
    headers = ["ID Agent", "Nom", "Prénom", "Poste",
               "Date", "Heure Arrivée", "Heure Départ", "Statut", "Observations"]
    for col, h in enumerate(headers, 1):
        cell = ws1.cell(row=4, column=col, value=h)
        style_header(cell, bg=BLEU_FONCE)
    ws1.row_dimensions[4].height = 22

    # Données agents (lignes vides pour saisie)
    jours = 26  # ~1 mois de travail
    row = 5
    fills = {
        "Présent": PatternFill("solid", fgColor=VERT),
        "Absent":  PatternFill("solid", fgColor=ROUGE),
        "Retard":  PatternFill("solid", fgColor=ORANGE),
    }
    for agent in agents:
        for j in range(jours):
            ws1.cell(row=row, column=1, value=agent["id"])
            ws1.cell(row=row, column=2, value=agent["nom"])
            ws1.cell(row=row, column=3, value=agent["prenom"])
            ws1.cell(row=row, column=4, value=agent["poste"])
            ws1.cell(row=row, column=5, value=f"Jour {j+1:02d}")
            ws1.cell(row=row, column=6, value="")  # heure arrivée
            ws1.cell(row=row, column=7, value="")  # heure départ
            ws1.cell(row=row, column=8, value="Présent")
            ws1.cell(row=row, column=8).fill = fills["Présent"]
            ws1.cell(row=row, column=9, value="")
            if row % 2 == 0:
                for c in range(1, 10):
                    if ws1.cell(row=row, column=c).fill.fgColor.rgb == "00000000":
                        ws1.cell(row=row, column=c).fill = PatternFill("solid", fgColor=GRIS)
            row += 1

    border_all(ws1, 4, row-1, 1, 9)

    # Largeurs colonnes
    widths = [10, 15, 15, 20, 10, 14, 14, 12, 22]
    for i, w in enumerate(widths, 1):
        ws1.column_dimensions[get_column_letter(i)].width = w

    ws1.freeze_panes = "A5"

    # ════════════════════════════════════════
    # FEUILLE 2 — RÉCAPITULATIF MENSUEL
    # ════════════════════════════════════════
    ws2 = wb.create_sheet("Récapitulatif Mensuel")
    ws2.sheet_view.showGridLines = False

    ws2.merge_cells("A1:H1")
    ws2["A1"] = "RÉCAPITULATIF MENSUEL DES PRÉSENCES"
    style_header(ws2["A1"], size=14)
    ws2.row_dimensions[1].height = 30

    headers2 = ["ID", "Nom", "Prénom", "Poste",
                "Jours Présents", "Jours Absents", "Retards", "Taux Présence (%)"]
    for col, h in enumerate(headers2, 1):
        cell = ws2.cell(row=3, column=col, value=h)
        style_header(cell, bg=BLEU_FONCE)
    ws2.row_dimensions[3].height = 22

    for i, agent in enumerate(agents, 1):
        r = 3 + i
        ws2.cell(row=r, column=1, value=agent["id"])
        ws2.cell(row=r, column=2, value=agent["nom"])
        ws2.cell(row=r, column=3, value=agent["prenom"])
        ws2.cell(row=r, column=4, value=agent["poste"])
        ws2.cell(row=r, column=5, value=0)
        ws2.cell(row=r, column=6, value=0)
        ws2.cell(row=r, column=7, value=0)
        tx = ws2.cell(row=r, column=8,
                      value=f"=IF((E{r}+F{r})=0,0,ROUND(E{r}/(E{r}+F{r})*100,1))")
        tx.number_format = '0.0"%"'
        if r % 2 == 0:
            for c in range(1, 9):
                ws2.cell(row=r, column=c).fill = PatternFill("solid", fgColor=GRIS)

    border_all(ws2, 3, 3 + len(agents), 1, 8)
    for i, w in enumerate([10, 15, 15, 20, 14, 14, 10, 18], 1):
        ws2.column_dimensions[get_column_letter(i)].width = w

    # ════════════════════════════════════════
    # FEUILLE 3 — REGISTRE AGENTS
    # ════════════════════════════════════════
    ws3 = wb.create_sheet("Registre Agents")
    ws3.sheet_view.showGridLines = False

    ws3.merge_cells("A1:G1")
    ws3["A1"] = "REGISTRE DES AGENTS — SERVICE QUALITÉ"
    style_header(ws3["A1"], size=14)
    ws3.row_dimensions[1].height = 30

    headers3 = ["ID Agent", "Nom", "Prénom", "Poste", "Date d'embauche", "QR Code (fichier)", "Statut"]
    for col, h in enumerate(headers3, 1):
        cell = ws3.cell(row=3, column=col, value=h)
        style_header(cell, bg=BLEU_FONCE)
    ws3.row_dimensions[3].height = 22

    for i, agent in enumerate(agents, 1):
        r = 3 + i
        ws3.cell(row=r, column=1, value=agent["id"])
        ws3.cell(row=r, column=2, value=agent["nom"])
        ws3.cell(row=r, column=3, value=agent["prenom"])
        ws3.cell(row=r, column=4, value=agent["poste"])
        ws3.cell(row=r, column=5, value=agent["date_embauche"])
        ws3.cell(row=r, column=6, value=f"QR_{agent['id']}_{agent['nom']}.png")
        ws3.cell(row=r, column=7, value="Actif")
        ws3.cell(row=r, column=7).fill = PatternFill("solid", fgColor=VERT)
        if r % 2 == 0:
            for c in range(1, 7):
                ws3.cell(row=r, column=c).fill = PatternFill("solid", fgColor=GRIS)

    border_all(ws3, 3, 3 + len(agents), 1, 7)
    for i, w in enumerate([12, 15, 15, 22, 16, 28, 10], 1):
        ws3.column_dimensions[get_column_letter(i)].width = w

    # ════════════════════════════════════════
    # FEUILLE 4 — STATISTIQUES
    # ════════════════════════════════════════
    ws4 = wb.create_sheet("Statistiques")
    ws4.sheet_view.showGridLines = False

    ws4.merge_cells("A1:D1")
    ws4["A1"] = "TABLEAU DE BORD — STATISTIQUES QUALITÉ"
    style_header(ws4["A1"], size=14)
    ws4.row_dimensions[1].height = 30

    stats = [
        ("", ""),
        ("INDICATEURS DE PRÉSENCE", ""),
        ("Total agents actifs", f"=COUNTA('Registre Agents'!A4:A1000)"),
        ("Jours ouvrés du mois", 26),
        ("", ""),
        ("RÈGLEMENT INTÉRIEUR", ""),
        ("Infractions téléphone (mois)", 0),
        ("Non-conformités vestimentaires", 0),
        ("Incidents bruit signalés", 0),
        ("Non-conformités hygiène", 0),
        ("", ""),
        ("DERNIÈRE MISE À JOUR", datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]

    for i, (label, val) in enumerate(stats, 3):
        ws4.cell(row=i, column=1, value=label)
        ws4.cell(row=i, column=2, value=val)
        if label and label == label.upper() and label != "":
            ws4.cell(row=i, column=1).font = Font(bold=True, color=BLEU_FONCE, size=11)
        elif label:
            ws4.cell(row=i, column=1).fill = PatternFill("solid", fgColor=BLEU_CLAIR)
            ws4.cell(row=i, column=2).fill = PatternFill("solid", fgColor=BLEU_CLAIR)

    ws4.column_dimensions["A"].width = 35
    ws4.column_dimensions["B"].width = 25

    out = EXPORT_DIR / "Gestion_Qualite_Pointage.xlsx"
    wb.save(str(out))
    print(f"  ✓ Fichier Excel : {out.name}")

# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == "__main__":
    print("\n=== Génération des documents Gestion Qualité ===\n")

    print("1. QR Codes agents...")
    generer_qrcodes()

    print("\n2. Règlement Intérieur (Word)...")
    generer_reglement_word()

    print("\n3. Fichier Excel (pointage + registre + stats)...")
    generer_excel_pointage()

    print("\n✅ Tous les documents ont été générés dans le dossier 'exports/' et 'qrcodes/'")
    print(f"   Répertoire : {BASE_DIR}")
